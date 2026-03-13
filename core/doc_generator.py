"""Word document generator - replaces VBA mWord.bas logic using python-docx."""

import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from core.models import ExtendedMarketData

logger = logging.getLogger(__name__)

# Formatting constants matching VBA
COLOR_DARK_BLUE = RGBColor(23, 52, 97)
COLOR_BLACK = RGBColor(0, 0, 0)


def _add_heading_text(doc: Document, text: str, size: int = 14,
                      bold: bool = True, color: RGBColor = COLOR_DARK_BLUE):
    """Add a formatted heading paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return para


def _add_body_text(doc: Document, text: str, italic: bool = False,
                   size: int = 11):
    """Add a body text paragraph."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    return para


def _add_bullet_item(doc: Document, text: str, level: int = 0):
    """Add a bulleted list item."""
    para = doc.add_paragraph(text, style="List Bullet")
    # Adjust indent for nested levels
    if level > 0:
        para.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return para


def _add_image_if_exists(doc: Document, images_dir: Path, image_name: str,
                         width: float = 5.5):
    """Add an image to the document if the file exists."""
    img_path = images_dir / image_name
    if img_path.exists():
        doc.add_picture(str(img_path), width=Inches(width))
        logger.info(f"Embedded image: {image_name}")
        return True
    else:
        logger.warning(f"Image not found: {img_path}")
        return False


def _build_scope_table(doc: Document, data: ExtendedMarketData):
    """Build the Market Report Scope table (replicates VBA CreateDocTable)."""
    table = doc.add_table(rows=10, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Market name (merged across all cols)
    row = table.rows[0]
    row.cells[0].merge(row.cells[3])
    row.cells[0].text = data.market_name

    # Row 1: Headers
    row = table.rows[1]
    row.cells[0].text = "Report Coverage"
    row.cells[1].text = "Details"

    # Row 2: Base Year + Market Size
    row = table.rows[2]
    row.cells[0].text = "Base Year"
    row.cells[1].text = str(data.base_year)
    row.cells[2].text = f"Market Size in {data.base_year}:"
    row.cells[3].text = data.market_size_value

    # Row 3: Historical + Forecast
    row = table.rows[3]
    row.cells[0].text = "Historical Data For:"
    row.cells[1].text = f"{data.historical_start} To {data.base_year - 1}"
    row.cells[2].text = "Forecast Period:"
    row.cells[3].text = f"{data.base_year} To {data.forecast_end}"

    # Row 4: CAGR + Projection
    row = table.rows[4]
    cagr_pct = f"{data.cagr_value * 100:.1f}" if data.cagr_value < 1 else f"{data.cagr_value:.1f}"
    row.cells[0].text = f"Forecast Period {data.base_year} To {data.forecast_end} CAGR:"
    row.cells[1].text = f"{cagr_pct}%"
    row.cells[2].text = f"{data.forecast_end} Value Projection:"
    row.cells[3].text = data.forecast_size_value

    # Row 5: Geographies
    row = table.rows[5]
    row.cells[0].text = "Geographies covered:"
    geo_lines = []
    for region, countries in data.geographic_data.items():
        geo_lines.append(f"{region}: {', '.join(countries)}")
    row.cells[1].merge(row.cells[3])
    row.cells[1].text = "\n".join(geo_lines)

    # Row 6: Segments
    row = table.rows[6]
    row.cells[0].text = "Segments covered:"
    seg_lines = []
    for seg in data.segments:
        seg_lines.append(f"{seg.name}: {', '.join(seg.sub_segments)}")
    row.cells[1].merge(row.cells[3])
    row.cells[1].text = "\n".join(seg_lines)

    # Row 7: Companies
    row = table.rows[7]
    row.cells[0].text = "Companies covered:"
    row.cells[1].merge(row.cells[3])
    row.cells[1].text = ", ".join(data.companies)

    # Row 8: Growth Drivers
    row = table.rows[8]
    row.cells[0].text = "Growth Drivers:"
    drivers_text = data.driver_1
    if data.driver_2:
        drivers_text += f"\n{data.driver_2}"
    row.cells[1].merge(row.cells[3])
    row.cells[1].text = drivers_text

    # Row 9: Restraints
    row = table.rows[9]
    row.cells[0].text = "Restraints & Challenges:"
    restraints_text = data.restraint_1
    if data.restraint_2:
        restraints_text += f"\n{data.restraint_2}"
    row.cells[1].merge(row.cells[3])
    row.cells[1].text = restraints_text

    # Bold the left-column labels
    for row in table.rows:
        for para in row.cells[0].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)

    return table


def _add_faqs(doc: Document, data: ExtendedMarketData):
    """Add FAQ section (replicates VBA CreateDocQuestionaries)."""
    _add_heading_text(doc, "Frequently Asked Questions")

    cagr_pct = f"{data.cagr_value * 100:.1f}" if data.cagr_value < 1 else f"{data.cagr_value:.1f}"

    faqs = [
        (
            f"What is the expected Compound Annual Growth Rate (CAGR) of the {data.market_name} during the forecast period ({data.base_year} - {data.forecast_end})?",
            f"The {data.market_name} is expected to grow at a CAGR of {cagr_pct}% during the forecast period ({data.base_year} - {data.forecast_end})."
        ),
        (
            f"What factors are driving the growth of the {data.market_name}?",
            f"Key factors driving the growth include: {data.driver_1}" +
            (f" and {data.driver_2}" if data.driver_2 else "") + "."
        ),
        (
            f"What are the factors hampering the growth of the {data.market_name}?",
            f"Factors hampering the growth include: {data.restraint_1}" +
            (f" and {data.restraint_2}" if data.restraint_2 else "") + "."
        ),
        (
            f"What is the leading {data.segments[0].name if data.segments else 'segment'} in the {data.market_name}?",
            f"The leading segment is {data.segments[0].dominating if data.segments else 'N/A'}."
        ),
        (
            f"Which are the major players operating in the {data.market_name}?",
            f"Major players include: {', '.join(data.companies[:5])}" +
            (f", and {len(data.companies) - 5} more" if len(data.companies) > 5 else "") + "."
        ),
        (
            f"Which region has the leading position in the {data.market_name}?",
            f"The {data.dominating_region} region holds the leading position in the market."
        ),
    ]

    for q, a in faqs:
        _add_heading_text(doc, f"Q: {q}", size=11, bold=True, color=COLOR_BLACK)
        _add_body_text(doc, f"A: {a}")


def _add_segmentation_section(doc: Document, data: ExtendedMarketData):
    """Add Market Segmentation section with multi-level bullets."""
    _add_heading_text(doc, "Market Segmentation")

    # Segment types with sub-segments
    for seg in data.segments:
        _add_heading_text(doc, seg.name, size=12, bold=True, color=COLOR_DARK_BLUE)
        for sub in seg.sub_segments:
            _add_bullet_item(doc, sub, level=0)

    # Regional breakdown
    _add_heading_text(doc, "By Region", size=12, bold=True, color=COLOR_DARK_BLUE)
    for region, countries in data.geographic_data.items():
        _add_bullet_item(doc, region, level=0)
        for country in countries:
            _add_bullet_item(doc, country, level=1)

    # Key Players
    _add_heading_text(doc, "Key Players", size=12, bold=True, color=COLOR_DARK_BLUE)
    for company in data.companies:
        _add_bullet_item(doc, company, level=0)


def generate_combined_doc(
    data: ExtendedMarketData,
    images_dir: Path,
    output_path: Path,
    rs_doc_path: Path = None,
) -> Path:
    """Generate the Combined document with embedded images and formatted sections.

    Replicates the VBA logic from mWord.bas for creating the main report document.

    Args:
        data: ExtendedMarketData with all fields populated
        images_dir: Directory containing the 5 JPG images
        output_path: Where to save the Combined .docx
        rs_doc_path: Optional path to RS .doc file for extracting section text

    Returns:
        Path to the saved document
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    images_dir = Path(images_dir)

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # 1. Title
    title_para = _add_heading_text(doc, f"Category - {data.market_name}",
                                   size=18, bold=True, color=COLOR_DARK_BLUE)

    # 2. Market Size and Trends
    _add_heading_text(doc, "Market Size and Trends")
    _add_body_text(doc, data.market_size_text)
    _add_image_if_exists(doc, images_dir, "Impact_Analysis.jpg")

    # 3. Market Driver
    if data.driver_1:
        _add_heading_text(doc, f"Market Driver - {data.driver_1[:80]}")
        _add_body_text(doc, data.driver_1)
    if data.driver_2:
        _add_heading_text(doc, f"Market Driver - {data.driver_2[:80]}")
        _add_body_text(doc, data.driver_2)

    # 4. Market Challenge
    if data.restraint_1:
        _add_heading_text(doc, f"Market Challenge - {data.restraint_1[:80]}")
        _add_body_text(doc, data.restraint_1)

    # 5. Segmental Analysis
    _add_heading_text(doc, "Segmental Analysis")
    _add_image_if_exists(doc, images_dir, "Segmental_Insights.jpg")
    for seg in data.segments:
        _add_body_text(doc, f"{seg.name}: The leading segment is {seg.dominating}.")

    # 6. Regional Insights
    _add_heading_text(doc, "Regional Insights")
    _add_image_if_exists(doc, images_dir, "Regional_Insights.jpg")
    if data.dominating_region:
        _add_body_text(doc, f"The {data.dominating_region} region dominates the market.")
    if data.fastest_growing_region:
        _add_body_text(doc, f"The {data.fastest_growing_region} region is the fastest growing.")

    # 7. Competitive Landscape
    _add_heading_text(doc, "Competitive Landscape")
    _add_image_if_exists(doc, images_dir, "Market_KeyPlayer.jpg")
    if data.companies:
        _add_body_text(doc, f"Key players include: {', '.join(data.companies)}.")

    # 8. Key Developments
    if data.takeaways:
        _add_heading_text(doc, "Key Developments")
        for takeaway in data.takeaways:
            _add_bullet_item(doc, takeaway)

    # 9. Key Takeaways from Analyst
    if data.takeaways:
        _add_heading_text(doc, "Key Takeaways from Analyst")
        for takeaway in data.takeaways:
            para = _add_body_text(doc, takeaway, italic=True)

    # 10. Market Report Scope table
    _add_heading_text(doc, "Market Report Scope")
    _build_scope_table(doc, data)

    # 11. FAQs
    _add_faqs(doc, data)

    # 12. Market Segmentation
    _add_segmentation_section(doc, data)

    doc.save(str(output_path))
    logger.info(f"Saved Combined document to {output_path}")
    return output_path


def generate_toc_doc(
    data: ExtendedMarketData,
    output_path: Path,
) -> Path:
    """Generate the Table of Contents document.

    Replicates VBA CreateTableOfContent from mWord.bas.

    Args:
        data: ExtendedMarketData with all fields populated
        output_path: Where to save the TOC .docx

    Returns:
        Path to the saved document
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    _add_heading_text(doc, f"{data.market_name} Report - Table of Contents",
                      size=16, bold=True, color=COLOR_DARK_BLUE)

    section_num = 1

    # Section 1: Research Objectives and Assumptions
    _add_heading_text(doc, f"Section {section_num}: Research Objectives and Assumptions",
                      size=12, bold=True)
    _add_bullet_item(doc, "Research Objectives")
    _add_bullet_item(doc, "Assumptions")
    _add_bullet_item(doc, "Abbreviations")
    section_num += 1

    # Section 2: Market Purview
    _add_heading_text(doc, f"Section {section_num}: Market Purview",
                      size=12, bold=True)
    _add_bullet_item(doc, "Report Description")
    _add_bullet_item(doc, "Market Definition and Scope")
    _add_bullet_item(doc, "Executive Summary")
    for seg in data.segments:
        _add_bullet_item(doc, f"Market Snippet, {seg.name}", level=1)
    _add_bullet_item(doc, "Market Snippet, By Region", level=1)
    _add_bullet_item(doc, "Coherent Opportunity Map (COM)")
    section_num += 1

    # Section 3: Market Dynamics, Regulations, and Trends Analysis
    _add_heading_text(doc, f"Section {section_num}: Market Dynamics, Regulations, and Trends Analysis",
                      size=12, bold=True)
    _add_bullet_item(doc, "Market Dynamics")
    _add_bullet_item(doc, "Drivers", level=1)
    _add_bullet_item(doc, "Restraints", level=1)
    _add_bullet_item(doc, "Market Opportunities", level=1)
    _add_bullet_item(doc, "Impact Analysis")
    _add_bullet_item(doc, "Key Highlights")
    _add_bullet_item(doc, "Regulatory Scenario")
    _add_bullet_item(doc, "Product Launch/Approvals")
    _add_bullet_item(doc, "PEST Analysis")
    _add_bullet_item(doc, "PORTER's Analysis")
    _add_bullet_item(doc, "Merger and Acquisition Scenario")
    section_num += 1

    # Dynamic segment sections
    for seg in data.segments:
        _add_heading_text(doc, f"Section {section_num}: {data.market_name}, {seg.name}",
                          size=12, bold=True)
        _add_bullet_item(doc, "Introduction")
        _add_bullet_item(doc, "Market Share Analysis")
        _add_bullet_item(doc, "Y-o-Y Growth Analysis")
        _add_bullet_item(doc, f"{seg.name} Trends")
        for sub in seg.sub_segments:
            _add_bullet_item(doc, sub, level=1)
            _add_bullet_item(doc, "Introduction", level=2)
            _add_bullet_item(doc, "Market Size and Forecast", level=2)
        section_num += 1

    # Regional section
    _add_heading_text(doc, f"Section {section_num}: {data.market_name}, By Region",
                      size=12, bold=True)
    _add_bullet_item(doc, "Introduction")
    _add_bullet_item(doc, "Market Share Analysis, By Region")
    _add_bullet_item(doc, "Y-o-Y Growth Analysis, By Region")
    for region, countries in data.geographic_data.items():
        _add_bullet_item(doc, region, level=1)
        _add_bullet_item(doc, "Introduction", level=2)
        _add_bullet_item(doc, "Market Size and Forecast", level=2)
        for country in countries:
            _add_bullet_item(doc, country, level=2)
    section_num += 1

    # Competitive Landscape
    _add_heading_text(doc, f"Section {section_num}: Competitive Landscape",
                      size=12, bold=True)
    _add_bullet_item(doc, "Market Share Analysis")
    for company in data.companies:
        _add_bullet_item(doc, company, level=1)
        _add_bullet_item(doc, "Company Highlights", level=2)
        _add_bullet_item(doc, "Product Portfolio", level=2)
        _add_bullet_item(doc, "Key Developments", level=2)
        _add_bullet_item(doc, "Financial Performance", level=2)
        _add_bullet_item(doc, "Strategies", level=2)
    section_num += 1

    # Analyst Recommendations
    _add_heading_text(doc, f"Section {section_num}: Analyst Recommendations",
                      size=12, bold=True)
    _add_bullet_item(doc, "Wheel of Fortune")
    _add_bullet_item(doc, "Analyst View")
    _add_bullet_item(doc, "Coherent Opportunity Map")
    section_num += 1

    # References
    _add_heading_text(doc, f"Section {section_num}: References and Methodology",
                      size=12, bold=True)
    _add_bullet_item(doc, "References")
    _add_bullet_item(doc, "Research Methodology")
    _add_bullet_item(doc, "About Us")

    doc.save(str(output_path))
    logger.info(f"Saved TOC document to {output_path}")
    return output_path
